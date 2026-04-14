import re
import datetime
import os
import time

import feedparser
import requests
from bs4 import BeautifulSoup
from deep_translator import GoogleTranslator
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 字体设置 ──────────────────────────────────────────────────────
EN_FONT = "Times New Roman"
CN_FONT = "SimSun"  # 宋体


def set_run_font(run, size_pt, bold=False, color=None, en_font=EN_FONT, cn_font=CN_FONT):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = en_font
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 设置东亚字体（中文）
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:cs"), cn_font)


def set_para_spacing(para, before_pt=0, after_pt=4, line_rule=None):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    if line_rule:
        spacing.set(qn("w:lineRule"), line_rule[0])
        spacing.set(qn("w:line"), line_rule[1])


# ── 新闻来源（RSS） ────────────────────────────────────────────────
FEEDS = {
    "🌍 国际政治 / World Politics": [
        "http://feeds.bbci.co.uk/news/world/rss.xml",
        "https://feeds.reuters.com/reuters/worldNews",
        "https://rss.nytimes.com/services/xml/rss/nyt/World.xml",
    ],
    "💰 经济财经 / Economy & Finance": [
        "https://feeds.reuters.com/reuters/businessNews",
        "https://feeds.bbci.co.uk/news/business/rss.xml",
        "https://rss.nytimes.com/services/xml/rss/nyt/Business.xml",
    ],
    "🔬 科技科学 / Technology & Science": [
        "https://feeds.bbci.co.uk/news/technology/rss.xml",
        "https://feeds.bbci.co.uk/news/science_and_environment/rss.xml",
        "https://rss.nytimes.com/services/xml/rss/nyt/Technology.xml",
    ],
    "🎭 文化社会 / Culture & Society": [
        "https://feeds.bbci.co.uk/news/entertainment_and_arts/rss.xml",
        "https://feeds.reuters.com/reuters/lifestyle",
        "https://rss.nytimes.com/services/xml/rss/nyt/Arts.xml",
    ],
    "🌿 气候环境 / Climate & Environment": [
        "https://www.theguardian.com/environment/rss",
        "https://feeds.bbci.co.uk/news/science_and_environment/rss.xml",
    ],
    "⚽ 体育 / Sports": [
        "https://feeds.bbci.co.uk/sport/rss.xml",
        "https://feeds.reuters.com/reuters/sportsNews",
    ],
}

ARTICLES_PER_SECTION = 5
SUMMARY_MAX_CHARS = 800
HEADERS = {"User-Agent": "Mozilla/5.0 (compatible; WeeklyNewsBot/1.0)"}

URL_SOURCE_MAP = {
    "bbci.co.uk": "BBC News",
    "bbc.co.uk": "BBC News",
    "reuters.com": "Reuters",
    "nytimes.com": "The New York Times",
    "theguardian.com": "The Guardian",
}


# ── 抓取正文 ──────────────────────────────────────────────────────
def fetch_full_text(url, fallback_summary):
    try:
        resp = requests.get(url, headers=HEADERS, timeout=8)
        soup = BeautifulSoup(resp.text, "lxml")
        # 尝试常见正文选择器
        for selector in ["article", "[class*='article-body']", "[class*='story-body']",
                         "main", "[class*='content']"]:
            container = soup.select_one(selector)
            if container:
                paras = container.find_all("p")
                text = " ".join(p.get_text(strip=True) for p in paras if len(p.get_text(strip=True)) > 40)
                if len(text) > 200:
                    return text[:SUMMARY_MAX_CHARS]
    except Exception:
        pass
    return fallback_summary[:SUMMARY_MAX_CHARS] if fallback_summary else ""


def clean_html(text):
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def source_from_url(url):
    for domain, name in URL_SOURCE_MAP.items():
        if domain in url:
            return name
    return ""


def fetch_articles(urls, limit):
    articles = []
    seen_titles = set()
    for url in urls:
        if len(articles) >= limit:
            break
        source = source_from_url(url)
        try:
            feed = feedparser.parse(url)
            for entry in feed.entries:
                if len(articles) >= limit:
                    break
                title = clean_html(entry.get("title", "")).strip()
                if not title or title in seen_titles:
                    continue
                seen_titles.add(title)
                rss_summary = clean_html(entry.get("summary", entry.get("description", "")))
                link = entry.get("link", "")
                full_text = fetch_full_text(link, rss_summary) if link else rss_summary[:SUMMARY_MAX_CHARS]
                # 来源优先用 RSS 自带的 source 字段
                entry_source = entry.get("source", {}).get("title", "") or source
                articles.append({"title": title, "summary": full_text, "link": link, "source": entry_source})
                time.sleep(0.3)
        except Exception as e:
            print(f"  [警告] 抓取 {url} 失败: {e}")
    return articles[:limit]


# ── 翻译 ──────────────────────────────────────────────────────────
def translate_to_chinese(text):
    if not text:
        return ""
    try:
        translator = GoogleTranslator(source="auto", target="zh-CN")
        chunks, result = [], []
        # 按句子分块，避免超出 5000 字符限制
        words = text.split(". ")
        chunk = ""
        for w in words:
            if len(chunk) + len(w) < 4000:
                chunk += w + ". "
            else:
                chunks.append(chunk.strip())
                chunk = w + ". "
        if chunk:
            chunks.append(chunk.strip())
        for c in chunks:
            result.append(translator.translate(c))
            time.sleep(0.2)
        return " ".join(result)
    except Exception as e:
        print(f"  [警告] 翻译失败: {e}")
        return ""


# ── 豆瓣书单抓取 ─────────────────────────────────────────────────
DOUBAN_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                  "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Referer": "https://book.douban.com/",
}

DOUBAN_CHART_URLS = [
    "https://book.douban.com/chart?subcat=G&time=week&type=rank",
    "https://book.douban.com/chart?subcat=G",
    "https://book.douban.com/chart",
]


def _douban_get(url, retries=3, backoff=2):
    """带退避重试的 requests.get，专用于豆瓣请求"""
    for attempt in range(retries):
        try:
            resp = requests.get(url, headers=DOUBAN_HEADERS, timeout=10)
            if resp.status_code == 200:
                return resp
            print(f"  [重试] {url} 返回 {resp.status_code}，第 {attempt + 1}/{retries} 次")
        except Exception as e:
            print(f"  [重试] {url} 异常: {e}，第 {attempt + 1}/{retries} 次")
        if attempt < retries - 1:
            time.sleep(backoff * (attempt + 1))
    return None


def fetch_douban_book_detail(url):
    """抓取豆瓣单本书详情页，返回 {author, category, rating, intro}"""
    info = {"author": "", "category": "", "rating": "", "intro": ""}
    try:
        resp = _douban_get(url)
        if resp is None:
            return info
        soup = BeautifulSoup(resp.text, "lxml")

        info_div = soup.select_one("#info")
        if info_div:
            text = info_div.get_text(" ", strip=True)
            # 提取作者
            author_match = re.search(r"作者[:：]\s*([^\n]+?)(?:\s*出版|$)", text)
            if author_match:
                info["author"] = author_match.group(1).strip()[:60]
            # 提取出版社/类别作为 fallback
            pub_match = re.search(r"出版社[:：]\s*([^\s]+)", text)
            if pub_match:
                info["category"] = pub_match.group(1).strip()

        # 豆瓣标签作为分类
        tags = soup.select(".tags-body a")
        if tags:
            info["category"] = " / ".join(t.get_text(strip=True) for t in tags[:4])

        # 评分
        rating_el = soup.select_one(".ll.rating_num")
        if rating_el:
            info["rating"] = rating_el.get_text(strip=True)

        # 简介
        intro_el = soup.select_one("#link-report .intro") or soup.select_one(".related_info .intro")
        if intro_el:
            paras = intro_el.find_all("p")
            intro_text = " ".join(p.get_text(strip=True) for p in paras if p.get_text(strip=True))
            info["intro"] = intro_text[:400] if intro_text else ""

        time.sleep(0.8)
    except Exception as e:
        print(f"    [警告] 书籍详情抓取失败 {url}: {e}")
    return info


def fetch_douban_weekly_books(top_n=5):
    """抓取豆瓣每周书榜 Top N"""
    books = []
    for chart_url in DOUBAN_CHART_URLS:
        try:
            resp = _douban_get(chart_url)
            if resp is None:
                continue
            soup = BeautifulSoup(resp.text, "lxml")

            # 尝试多种选择器
            items = (
                soup.select(".media.clearfix") or
                soup.select("li.media") or
                soup.select(".chart-dashed-list li")
            )
            if not items:
                continue

            for item in items[:top_n]:
                title_el = item.select_one("h2 a") or item.select_one("a[title]")
                if not title_el:
                    continue
                title = title_el.get_text(strip=True)
                link = title_el.get("href", "")

                # 尝试从列表页获取基础信息
                meta_el = item.select_one(".color-gray") or item.select_one("p.author")
                meta_text = meta_el.get_text(" ", strip=True) if meta_el else ""

                print(f"    📖 {title[:40]}...")
                detail = fetch_douban_book_detail(link) if link else {}

                books.append({
                    "title": title,
                    "link": link,
                    "author": detail.get("author") or meta_text[:60],
                    "category": detail.get("category", ""),
                    "rating": detail.get("rating", ""),
                    "intro": detail.get("intro", ""),
                })

            if books:
                break  # 成功拿到数据就不继续尝试
        except Exception as e:
            print(f"  [警告] 豆瓣榜单抓取失败 {chart_url}: {e}")

    return books


# ── 生成 Word 文档 ────────────────────────────────────────────────
def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_spacing(p, before_pt=0, after_pt=0)


def render_book_section(doc, books):
    """在文档末尾渲染豆瓣书单板块"""
    sec_p = doc.add_paragraph()
    set_para_spacing(sec_p, before_pt=18, after_pt=8)
    sec_run = sec_p.add_run("📚 每周阅读推荐 · 豆瓣书榜 Top 10")
    set_run_font(sec_run, size_pt=14, bold=True, color=(0x70, 0x30, 0xA0))

    if not books:
        np = doc.add_paragraph()
        r = np.add_run("（本周暂无数据，请稍后查看豆瓣读书）")
        set_run_font(r, size_pt=10, color=(0x99, 0x99, 0x99))
        add_horizontal_rule(doc)
        return

    for i, book in enumerate(books, 1):
        # 书名
        title_p = doc.add_paragraph()
        set_para_spacing(title_p, before_pt=12, after_pt=3)
        t_run = title_p.add_run(f"{i}.  {book['title']}")
        set_run_font(t_run, size_pt=14, bold=True, color=(0x1A, 0x1A, 0x1A),
                     en_font="Times New Roman", cn_font="SimSun")

        # 元信息行：作者 / 分类 / 评分
        meta_parts = []
        if book.get("author"):
            meta_parts.append(f"作者：{book['author']}")
        if book.get("category"):
            meta_parts.append(f"分类：{book['category']}")
        if book.get("rating"):
            meta_parts.append(f"豆瓣评分：{book['rating']}")
        if meta_parts:
            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.left_indent = Cm(0.8)
            set_para_spacing(meta_p, before_pt=0, after_pt=4)
            m_run = meta_p.add_run("  ·  ".join(meta_parts))
            set_run_font(m_run, size_pt=9, color=(0x66, 0x66, 0x66))

        # 内容简介
        if book.get("intro"):
            intro_p = doc.add_paragraph()
            intro_p.paragraph_format.left_indent = Cm(0.8)
            set_para_spacing(intro_p, before_pt=0, after_pt=8)
            i_run = intro_p.add_run(book["intro"])
            set_run_font(i_run, size_pt=10, color=(0x33, 0x33, 0x33))

    add_horizontal_rule(doc)


def build_document(sections_data, week_str, date_str, books=None):
    doc = Document()

    # 页面边距
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(3)

    # ── 大标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(title_p, before_pt=0, after_pt=6)
    r = title_p.add_run("每周全球时事新闻简报")
    set_run_font(r, size_pt=24, bold=True, color=(0x1F, 0x49, 0x7D))

    # 英文副标题
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(sub_p, before_pt=0, after_pt=2)
    r2 = sub_p.add_run("Weekly Global News Briefing")
    set_run_font(r2, size_pt=13, color=(0x5A, 0x5A, 0x5A))

    # 日期行
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(date_p, before_pt=0, after_pt=16)
    r3 = date_p.add_run(week_str)
    set_run_font(r3, size_pt=10, color=(0x99, 0x99, 0x99))

    add_horizontal_rule(doc)

    for section_name, articles in sections_data.items():
        # 板块标题
        sec_p = doc.add_paragraph()
        set_para_spacing(sec_p, before_pt=18, after_pt=8)
        sec_run = sec_p.add_run(section_name)
        set_run_font(sec_run, size_pt=14, bold=True, color=(0x2E, 0x74, 0xB5))

        if not articles:
            np = doc.add_paragraph()
            r = np.add_run("（本周暂无相关新闻 / No articles available this week）")
            set_run_font(r, size_pt=10, color=(0x99, 0x99, 0x99))
            continue

        for i, art in enumerate(articles, 1):
            # 英文标题（Times New Roman，字号加大）
            en_p = doc.add_paragraph()
            set_para_spacing(en_p, before_pt=12, after_pt=2)
            en_run = en_p.add_run(f"{i}.  {art['en_title']}")
            set_run_font(en_run, size_pt=13, bold=True, color=(0x1A, 0x1A, 0x1A),
                         en_font="Times New Roman", cn_font="SimSun")

            # 中文标题（宋体，字号加大）
            cn_p = doc.add_paragraph()
            set_para_spacing(cn_p, before_pt=0, after_pt=6)
            cn_run = cn_p.add_run(f"     {art['cn_title']}")
            set_run_font(cn_run, size_pt=13, bold=True, color=(0x2E, 0x74, 0xB5),
                         en_font="Times New Roman", cn_font="SimSun")

            # 英文正文（Times New Roman）
            if art["en_summary"]:
                en_body = doc.add_paragraph()
                en_body.paragraph_format.left_indent = Cm(0.8)
                set_para_spacing(en_body, before_pt=0, after_pt=4)
                r = en_body.add_run(art["en_summary"])
                set_run_font(r, size_pt=10, color=(0x33, 0x33, 0x33),
                             en_font="Times New Roman", cn_font="SimSun")

            # 中文正文（宋体）
            if art["cn_summary"]:
                cn_body = doc.add_paragraph()
                cn_body.paragraph_format.left_indent = Cm(0.8)
                set_para_spacing(cn_body, before_pt=0, after_pt=6)
                r = cn_body.add_run(art["cn_summary"])
                set_run_font(r, size_pt=10, color=(0x33, 0x33, 0x33),
                             en_font="Times New Roman", cn_font="SimSun")

            # 来源标注
            if art.get("source"):
                src_p = doc.add_paragraph()
                src_p.paragraph_format.left_indent = Cm(0.8)
                set_para_spacing(src_p, before_pt=0, after_pt=10)
                src_run = src_p.add_run(f"来源：{art['source']}")
                set_run_font(src_run, size_pt=8, color=(0xAA, 0xAA, 0xAA),
                             en_font="Times New Roman", cn_font="SimSun")

        add_horizontal_rule(doc)

    # 书单板块
    render_book_section(doc, books or [])

    # 页脚
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(footer_p, before_pt=12, after_pt=0)
    r = footer_p.add_run(
        f"来源：BBC News · Reuters · NYT · The Guardian · 豆瓣读书  ·  生成日期：{date_str}"
    )
    set_run_font(r, size_pt=8, color=(0xAA, 0xAA, 0xAA))

    return doc


# ── 主流程 ────────────────────────────────────────────────────────
def main():
    today = datetime.date.today()
    week_str = today.strftime("%Y 年第 %W 周  ·  Week %W of %Y  ·  %B %d, %Y")
    date_str = today.strftime("%Y-%m-%d")
    filename = f"WeeklyNews_{date_str}.docx"
    output_dir = "reports"
    os.makedirs(output_dir, exist_ok=True)

    print(f"📰 开始抓取新闻 ({today}) ...")
    sections_data = {}

    for section_name, urls in FEEDS.items():
        print(f"\n  ▶ {section_name}")
        raw_articles = fetch_articles(urls, ARTICLES_PER_SECTION)
        processed = []
        for art in raw_articles:
            print(f"    翻译标题: {art['title'][:60]}...")
            cn_title = translate_to_chinese(art["title"])
            cn_summary = translate_to_chinese(art["summary"]) if art["summary"] else ""
            processed.append({
                "en_title": art["title"],
                "cn_title": cn_title,
                "en_summary": art["summary"],
                "cn_summary": cn_summary,
                "source": art.get("source", ""),
            })
        sections_data[section_name] = processed

    print(f"\n📚 抓取豆瓣每周书榜...")
    books = fetch_douban_weekly_books(top_n=10)
    print(f"  共获取 {len(books)} 本书")

    print(f"\n📄 生成 Word 文档...")
    doc = build_document(sections_data, week_str, date_str, books=books)
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    print(f"✅ 已保存: {output_path}")


if __name__ == "__main__":
    main()
